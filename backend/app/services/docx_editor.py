import docx
from docx.shared import Pt
from docx.enum.shape import WD_INLINE_SHAPE_TYPE
import logging
import re

logger = logging.getLogger(__name__)

def get_docx_images_count(doc_path: str) -> int:
    try:
        doc = docx.Document(doc_path)
        count = 0
        for shape in doc.inline_shapes:
            if shape.type == WD_INLINE_SHAPE_TYPE.PICTURE:
                count += 1
        return count
    except Exception as e:
        logger.error("Error counting images in DOCX: %s", e)
        return 0

def apply_docx_styling(
    doc_path: str,
    output_path: str,
    font_name: str = None,
    font_size: float = None,
    header_text: str = None,
    footer_text: str = None,
    image_replacements: list = None  # List of dicts: [{"target_index": int, "image_bytes": bytes}]
) -> bool:
    try:
        doc = docx.Document(doc_path)

        # 1. Update font and sizes for paragraphs
        if font_name or font_size:
            # Change font of regular paragraphs
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if font_name:
                        run.font.name = font_name
                    if font_size:
                        run.font.size = Pt(font_size)

            # Change font inside tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if font_name:
                                    run.font.name = font_name
                                if font_size:
                                    run.font.size = Pt(font_size)

        # 2. Update Header/Footer
        for section in doc.sections:
            if header_text is not None:
                header = section.header
                # Clear existing header paragraphs
                for p in header.paragraphs:
                    p.text = ""
                if not header.paragraphs:
                    header.add_paragraph(header_text)
                else:
                    header.paragraphs[0].text = header_text

            if footer_text is not None:
                footer = section.footer
                # Clear existing footer paragraphs
                for p in footer.paragraphs:
                    p.text = ""
                if not footer.paragraphs:
                    footer.add_paragraph(footer_text)
                else:
                    footer.paragraphs[0].text = footer_text

        # 3. Image replacement
        if image_replacements:
            # Collect all shapes that are pictures
            pics = []
            for shape in doc.inline_shapes:
                if shape.type == WD_INLINE_SHAPE_TYPE.PICTURE:
                    pics.append(shape)

            for rep in image_replacements:
                idx = rep.get("target_index")
                img_bytes = rep.get("image_bytes")
                if idx is not None and img_bytes and 0 <= idx < len(pics):
                    target_shape = pics[idx]
                    try:
                        # Access internal XML element for embedding
                        rId = target_shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                        # Update binary blob in Zip archive package
                        image_part = doc.part.related_parts[rId]
                        image_part._blob = img_bytes
                        logger.info("Successfully replaced DOCX image at index %d", idx)
                    except Exception as img_err:
                        logger.error("Failed to replace image at index %d: %s", idx, img_err)

        doc.save(output_path)
        return True
    except Exception as e:
        logger.error("Error applying DOCX styling: %s", e)
        return False

def get_docx_content(doc_path: str) -> list:
    try:
        doc = docx.Document(doc_path)
        content = []
        for i, p in enumerate(doc.paragraphs):
            # Only send paragraphs with text, but allow empty ones in list for tracking
            content.append({
                "index": i,
                "text": p.text
            })
        return content
    except Exception as e:
        logger.error("Error getting DOCX content: %s", e)
        return []

def update_docx_content(doc_path: str, output_path: str, edits: list) -> bool:
    try:
        doc = docx.Document(doc_path)
        # edits is list of {"index": int, "text": str}
        edit_map = {edit["index"]: edit["text"] for edit in edits}
        
        for idx, p in enumerate(doc.paragraphs):
            if idx in edit_map:
                new_text = edit_map[idx]
                if p.text != new_text:
                    if p.runs:
                        # Update first run text to preserve styling and clear subsequent runs
                        p.runs[0].text = new_text
                        for run in p.runs[1:]:
                            run.text = ""
                    else:
                        p.text = new_text
        
        doc.save(output_path)
        return True
    except Exception as e:
        logger.error("Error updating DOCX content: %s", e)
        return False

def find_replace_docx(doc_path: str, output_path: str, find_text: str, replace_text: str, case_sensitive: bool = True) -> int:
    try:
        doc = docx.Document(doc_path)
        count = 0
        flags = 0 if case_sensitive else re.IGNORECASE
        pattern = re.compile(re.escape(find_text), flags)

        # 1. Replace in body paragraphs
        for p in doc.paragraphs:
            if pattern.search(p.text):
                new_text, n = pattern.subn(replace_text, p.text)
                if n > 0:
                    count += n
                    if p.runs:
                        p.runs[0].text = new_text
                        for run in p.runs[1:]:
                            run.text = ""
                    else:
                        p.text = new_text

        # 2. Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if pattern.search(p.text):
                            new_text, n = pattern.subn(replace_text, p.text)
                            if n > 0:
                                count += n
                                if p.runs:
                                    p.runs[0].text = new_text
                                    for run in p.runs[1:]:
                                        run.text = ""
                                else:
                                    p.text = new_text

        doc.save(output_path)
        return count
    except Exception as e:
        logger.error("Error in find-replace DOCX: %s", e)
        return 0
